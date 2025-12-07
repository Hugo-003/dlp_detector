# pip install PyPDF2 python-docx reportlab

import re
import os
import datetime
from PyPDF2 import PdfReader
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# =========================
# CONFIGURACIÓN
# =========================

MASKED_DIR = "masked_files"  # Carpeta donde vamos a meter todos los archivos enmascarados, para no liarnos con los originales
LOG_FILE = "dlp_logs.txt"    # Archivo donde vamos a ir guardando todos los logs de lo que detectamos

# Patrones de datos sensibles que queremos detectar
PATTERNS = {
    "DNI": r"\b\d{8}[A-HJ-NP-TV-Z]\b",
    "EMAIL": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",
    "TELEFONO": r"\b(?:\+34)?[6-9]\d{8}\b",
    "IBAN": r"\bES\d{22}\b",
    "TARJETA": r"\b(?:\d[ -]*?){13,16}\b"
}

# Severidad asociada a cada tipo de dato, para saber cuáles son más críticos
SEVERITY = {
    "DNI": "ALTO",
    "EMAIL": "MEDIO",
    "TELEFONO": "ALTO",
    "IBAN": "CRITICO",
    "TARJETA": "CRITICO"
}

# =========================
# FUNCIONES
# =========================

def ensure_masked_dir(base_path):
    # Crear carpeta para archivos enmascarados si no existe
    masked_path = os.path.join(base_path, MASKED_DIR)
    os.makedirs(masked_path, exist_ok=True)
    return masked_path

def log_event(event):
    # Guardamos los eventos en un log, para poder auditar después si hace falta
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(f"[{timestamp}] {event}\n")

def mask_data(data, label):
    # Reemplazamos los datos sensibles por algo más seguro, sin perder la referencia mínima
    if label in ["DNI", "TARJETA", "IBAN"]:
        return "*" * (len(data) - 2) + data[-2:]
    elif label == "EMAIL":
        parts = data.split("@")
        return parts[0][0] + "***@" + parts[1]
    elif label == "TELEFONO":
        return "***" + data[-3:]
    return data

def read_file(file_path):
    # Intentamos leer el archivo según su tipo, siempre controlando errores para que no crashee
    if file_path.lower().endswith(('.txt', '.csv', '.log')):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif file_path.lower().endswith('.pdf'):
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        except:
            # Si falla al leer el PDF, no paramos el programa
            pass
        return text
    elif file_path.lower().endswith('.docx'):
        text = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except:
            # Lo mismo con Word, seguimos aunque haya algún fallo
            pass
        return text
    return ""

def save_masked(content, original_path, masked_dir):
    # Guardamos la versión enmascarada, manteniendo el tipo de archivo
    base, ext = os.path.splitext(os.path.basename(original_path))
    if ext.lower() in ['.txt', '.csv', '.log']:
        masked_file = os.path.join(masked_dir, f"masked_{base}{ext}")
        with open(masked_file, "w", encoding="utf-8") as f:
            f.write(content)
    elif ext.lower() == '.pdf':
        masked_file = os.path.join(masked_dir, f"masked_{base}.pdf")
        c = canvas.Canvas(masked_file, pagesize=letter)
        y = 750
        for line in content.split("\n"):
            c.drawString(30, y, line[:90])
            y -= 15
            if y < 50:
                c.showPage()
                y = 750
        c.save()
    elif ext.lower() == '.docx':
        masked_file = os.path.join(masked_dir, f"masked_{base}.docx")
        doc = Document()
        for line in content.split("\n"):
            doc.add_paragraph(line)
        doc.save(masked_file)
    return masked_file

def scan_file(file_path, masked_dir):
    if not os.path.exists(file_path):
        return {"error": "Archivo no encontrado"}

    content = read_file(file_path)
    if not content:
        return {"error": "No se pudo leer el archivo o está vacío"}

    findings = {}
    masked_content = content
    total_incidents = 0
    severity_counter = {}

    for label, pattern in PATTERNS.items():
        # Buscamos coincidencias y vamos contando
        matches = re.findall(pattern, content)
        if matches:
            findings[label] = matches
            severity_counter[label] = len(matches)
            for match in matches:
                masked = mask_data(match, label)
                masked_content = masked_content.replace(match, masked)
                total_incidents += 1

    # Guardamos el archivo enmascarado en la carpeta dedicada
    if findings:
        masked_file = save_masked(masked_content, file_path, masked_dir)
    else:
        masked_file = None

    # Registramos en el log
    if findings:
        log_event(f"ALERTA: {file_path} | Tipos: {list(findings.keys())} | Incidentes: {total_incidents}")
    else:
        log_event(f"OK: Archivo limpio -> {file_path}")

    return {
        "findings": findings,
        "total": total_incidents,
        "severity": severity_counter,
        "masked_file": masked_file
    }

def scan_folder(folder_path):
    if not os.path.exists(folder_path):
        print("Carpeta no encontrada")
        return

    masked_dir = ensure_masked_dir(folder_path)

    all_results = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith(('.txt', '.csv', '.log', '.pdf', '.docx')):
                path = os.path.join(root, file)
                results = scan_file(path, masked_dir)
                results["file"] = path
                all_results.append(results)

    # Resumen global para el usuario
    total_files = len(all_results)
    files_with_incidents = sum(1 for r in all_results if r["total"] > 0)
    total_incidents = sum(r["total"] for r in all_results)

    print("\n=== RESUMEN GLOBAL ===")
    print(f"Archivos analizados: {total_files}")
    print(f"Archivos con incidentes: {files_with_incidents}")
    print(f"Total de incidentes: {total_incidents}\n")
    print(f"Todos los archivos enmascarados están en: {masked_dir}\n")

# =========================
# EJECUCIÓN
# =========================

if __name__ == "__main__":
    print("\n=== SISTEMA DLP AVANZADO (Archivos en carpeta masked_files) ===\n")
    raw_path = input("Introduce la ruta del archivo o carpeta a analizar: ").strip()
    # Limpiamos comillas o caracteres extraños que pueda meter el usuario
    raw_path = raw_path.strip('"').strip("'")
    path = os.path.normpath(raw_path)

    if os.path.isfile(path):
        base_dir = os.path.dirname(path)
        masked_dir = ensure_masked_dir(base_dir)
        results = scan_file(path, masked_dir)
        if "error" in results:
            print("Error:", results["error"])
        else:
            if results["findings"]:
                print("DATOS SENSIBLES ENCONTRADOS:\n")
                for k, v in results["findings"].items():
                    print(f"{k}: {len(v)} coincidencias | Severidad: {SEVERITY[k]}")
                print(f"\nTotal incidentes: {results['total']}")
                print(f"Archivo enmascarado generado en: {results['masked_file']}")
            else:
                print("Archivo limpio")
    elif os.path.isdir(path):
        scan_folder(path)
    else:
        print("Ruta no válida")
